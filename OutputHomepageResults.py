import json
import pandas as pd
from docx import Document
from docx.shared import Pt


with open("homepage_data_analyzed.json", "r") as f:
    data = json.load(f)
    analysis_df = pd.DataFrame(data)

    df = pd.read_csv("homepage_data.csv")
    df2 = pd.read_csv("homepage_data_3.csv")
    df = df[~df['Personal Website'].str.endswith('.sg').astype(bool)]
    df = pd.concat([df, df2], ignore_index=True)

    print("Analysis df columns: ", analysis_df.columns)
    print("df columns: ", df.columns)

    df = df.merge(analysis_df, left_on='Name', right_on='name', how='left')
    df = df[:5]

    with open("homepage_results_formatted.txt", "w") as output_file:
        doc = Document()
        doc.add_heading('Homepage Results', 0)
        for index, row in df.iterrows():
            doc.add_heading(f"Name: {row['Name']}", level=1)
            for col in df.columns:
                if "name" in col.lower():
                    continue
                doc.add_heading(f"{col}:", level=2)
                p = doc.add_paragraph(f"{row[col]}\n")
        doc.save("homepage_results_formatted.docx")